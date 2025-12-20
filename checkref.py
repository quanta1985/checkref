import streamlit as st
import re
import time
import pandas as pd
from docx import Document
from pypdf import PdfReader
from thefuzz import fuzz # Thư viện AI

# --- 1. CẤU HÌNH & CSS (GIỮ NGUYÊN 100%) ---
st.set_page_config(
    page_title="Citation Pro | AI Fuzzy Logic",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    /* Font và màu nền tổng thể */
    .stApp { background-color: #f8f9fa; }
    
    /* Style cho các Card (Khối) */
    .css-card {
        border-radius: 15px; padding: 20px; background-color: white;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05); margin-bottom: 20px; border: 1px solid #e9ecef;
    }
    
    /* Header chính */
    .main-header { font-family: 'Helvetica Neue', sans-serif; color: #0d6efd; text-align: center; margin-bottom: 30px; }
    
    /* Metric Box */
    div[data-testid="stMetric"] {
        background-color: #ffffff; border: 1px solid #e0e0e0; padding: 15px;
        border-radius: 10px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); text-align: center;
    }
    
    /* Alert Boxes */
    .alert-error { padding: 12px; border-radius: 8px; background-color: #fff5f5; border-left: 5px solid #fc8181; color: #c53030; margin-bottom: 10px; font-size: 15px; }
    .alert-warning { padding: 12px; border-radius: 8px; background-color: #fffaf0; border-left: 5px solid #f6ad55; color: #c05621; margin-bottom: 10px; font-size: 15px; }
    .alert-success { padding: 12px; border-radius: 8px; background-color: #f0fff4; border-left: 5px solid #48bb78; color: #2f855a; font-weight: bold; }
    .beta-note { font-size: 13px; color: #6c757d; font-style: italic; text-align: center; margin-bottom: 20px; }
</style>
""", unsafe_allow_html=True)

# --- 2. CÁC HÀM XỬ LÝ (LOGIC v12.2 - Data Cleaner) ---

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs: full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs: full_text.append(para.text)
        return "\n".join(full_text)
    except: return "ERROR_DOC"

def extract_text_from_pdf(file):
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages: text += page.extract_text() + "\n"
        return text
    except: return "ERROR_PDF"

def preprocess_text(text):
    text = re.sub(r'-\s*\n\s*', '', text)
    text = text.replace('\n', ' ').replace('\r', ' ')
    text = re.sub(r'\s+', ' ', text)
    return text

def is_legal_or_standard(text):
    text_lower = text.lower()
    keywords = [
        'tcvn', 'qcvn', 'iso', 'luật', 'nghị định', 'quyết định', 'thông tư', 
        'chỉ thị', 'qđ-ttg', 'nd-cp', 'tt-btnmt', 'luat', 'nghi dinh', 
        'quyet dinh', 'thong tu', 'tiêu chuẩn', 'quy chuẩn', 'chính phủ', 
        'quốc hội', 'bộ tài nguyên', 'bộ xây dựng', 'bộ khoa học', 'bộ tnmt'
    ]
    for kw in keywords:
        if kw in text_lower: return True
    return False

def is_garbage(text):
    text_lower = text.lower()
    # Cập nhật danh sách từ khóa rác (data keywords)
    blacklist = [
        'tháng', 'ngày', 'năm', 'lúc', 'trước', 'sau', 'khoảng', 'hình', 'bảng', 'biểu', 
        'sơ đồ', 'phương trình', 'công thức', 'hệ số', 'giá trị', 'tỉ lệ', 'kết quả', 
        'đoạn', 'phần', 'mục', 'bản đồ', 'giai đoạn', 'số', 'nghiên cứu', 'phân tích', 
        'đánh giá', 'đối với', 'của', 'bởi', 'được', 'trong', 'tại', 
        'tương đương', 'tương đương với', 'dao động', 'đến', 'từ' # <--- Thêm từ khóa mới
    ]
    for word in blacklist:
        if re.search(r'\b' + re.escape(word) + r'\b', text_lower):
            return True
    
    invalid_chars = ['/', '=', '>', '<', '%', '+', '\\']
    for char in invalid_chars:
        if char in text: return True
        
    return False

def expand_abbreviation(name):
    abbr_dict = {
        'BỘ TNMT': 'BỘ TÀI NGUYÊN VÀ MÔI TRƯỜNG',
        'BỘ TN&MT': 'BỘ TÀI NGUYÊN VÀ MÔI TRƯỜNG',
        'BTNMT': 'BỘ TÀI NGUYÊN VÀ MÔI TRƯỜNG',
        'BỘ KHCN': 'BỘ KHOA HỌC VÀ CÔNG NGHỆ',
        'BỘ NNPTNT': 'BỘ NÔNG NGHIỆP VÀ PHÁT TRIỂN NÔNG THÔN',
        'BỘ XD': 'BỘ XÂY DỰNG',
        'UBND': 'ỦY BAN NHÂN DÂN',
        'CP': 'CHÍNH PHỦ',
        'QH': 'QUỐC HỘI'
    }
    
    name_upper = name.upper()
    for abbr, full in abbr_dict.items():
        if abbr in name_upper:
            return name_upper.replace(abbr, full)
    return name

def check_citation_fuzzy(cit_name, cit_year, refs_list, threshold=65):
    if is_legal_or_standard(cit_name): return True

    clean_cit = re.sub(r'(et\s*al\.?|và\s*nnk\.?|và\s*cộng\s*sự|và\s*cs\.?|&\s*cs\.?|&|and)', ' ', cit_name, flags=re.IGNORECASE).strip()
    clean_cit = re.sub(r'^(được|bởi|của|theo)\s+', '', clean_cit, flags=re.IGNORECASE).strip()
    
    expanded_cit = expand_abbreviation(clean_cit)

    for ref in refs_list:
        if str(cit_year) in ref:
            # 1. Prefix Check
            clean_ref_start = re.sub(r'^\s*(\[?\d+\]?\.?)\s+', '', ref, count=1)
            if clean_ref_start.lower().startswith(clean_cit.lower()):
                return True
            
            # 2. Author Isolation
            match_year = re.search(str(cit_year), ref)
            if match_year:
                author_part_only = ref[:match_year.start()]
                score_author = fuzz.token_set_ratio(clean_cit, author_part_only)
                if score_author >= threshold: return True

            # 3. Full String Backup
            score1 = fuzz.token_set_ratio(clean_cit, ref)
            score2 = 0
            if expanded_cit != clean_cit:
                score2 = fuzz.token_set_ratio(expanded_cit, ref)
            
            if max(score1, score2) >= threshold:
                return True
    return False

def find_citations_v12(text):
    citations = []
    
    # Pattern 1: Trong ngoặc (...)
    for match in re.finditer(r'\(([^)]*?\d{4}[^)]*?)\)', text):
        content = match.group(1)
        for part in content.split(';'):
            part = part.strip()
            year_match = re.search(r'(\d{4})[a-z]?', part) 
            if year_match:
                year = year_match.group(1)
                name_part = part[:year_match.start()].strip().rstrip(',:').strip()
                
                # === NEW FIX v12.2: Tên tác giả KHÔNG ĐƯỢC chứa số ===
                # Nếu name_part chứa bất kỳ chữ số nào (0-9) -> Skip ngay lập tức
                # VD: "tương đương với 560,79" -> Có số -> Loại
                if re.search(r'\d', name_part):
                    continue
                # ====================================================

                if len(name_part) > 1 and len(name_part) < 100 and not is_legal_or_standard(name_part):
                     if not is_garbage(name_part):
                        citations.append({"name": name_part, "year": year, "full": f"({name_part}, {year})"})

    # Pattern 2: Dạng mở Name (Year)
    for match in re.finditer(r'([A-ZÀ-ỹ][A-Za-zÀ-ỹ\s&\-,]{1,60}?)\s*\(\s*(\d{4})\s*\)', text):
        raw_name = match.group(1).strip()
        year = match.group(2)
        
        # Áp dụng logic tương tự: Tên không được chứa số
        if re.search(r'\d', raw_name): continue

        if not is_legal_or_standard(raw_name) and not is_garbage(raw_name):
             citations.append({"name": raw_name, "year": year, "full": f"{raw_name} ({year})"})

    unique_citations = []
    seen = set()
    for c in citations:
        key = f"{c['name']}_{c['year']}"
        if key not in seen:
            unique_citations.append(c)
            seen.add(key)
    return unique_citations

# --- 3. GIAO DIỆN CHÍNH (GIỮ NGUYÊN) ---

# --- SIDEBAR ---
with st.sidebar:
    st.markdown("<h2 style='text-align: center; color: #0d6efd;'>🎓 Citation Pro <br><span style='font-size:16px; color: #666;'>(AI FUZZY CHECK)</span></h2>", unsafe_allow_html=True)
    st.markdown("---")
    uploaded_file = st.file_uploader("📂 **Tải báo cáo lên đây**:", type=['docx', 'pdf'])
    
    st.markdown("---")
    with st.expander("ℹ️ Hướng dẫn sử dụng"):
        st.markdown("""
        1. Upload file báo cáo (.docx/.pdf).
        2. Chờ hệ thống tự động quét.
        3. Xem kết quả tại Dashboard bên phải.
        """)
    
    st.info("⚠️ **Lưu ý:** App đang trong quá trình phát triển (Beta). Kết quả kiểm tra chỉ mang tính chất tham khảo nhanh.")
    st.caption("Dev by Quan HUMG")

# --- MAIN PAGE ---
if not uploaded_file:
    st.markdown("<div style='text-align: center; padding: 50px;'>", unsafe_allow_html=True)
    st.title("Công cụ Rà soát Trích dẫn & Tài liệu tham khảo")
    st.markdown("### 🚀 Nhanh chóng - (Gần) Chính xác - (Sắp) Chuyên nghiệp - Và JUST FOR FUN 😎")
    st.markdown("Kiểm tra sự đồng bộ giữa *Trích dẫn trong bài (In-text)* và *Danh mục tham khảo (References)*.")
    st.image("https://cdn-icons-png.flaticon.com/512/8662/8662266.png", width=150)
    st.info("👈 Vui lòng tải file báo cáo ở thanh bên trái để bắt đầu.")
    st.markdown("</div>", unsafe_allow_html=True)

else:
    # --- XỬ LÝ DỮ LIỆU ---
    with st.container():
        with st.status("Đang phân tích dữ liệu...", expanded=True) as status:
            time.sleep(0.3)
            st.write("📄 Đang đọc và làm sạch file...")
            
            if uploaded_file.name.endswith('.docx'):
                raw_text = extract_text_from_docx(uploaded_file)
            else:
                raw_text = extract_text_from_pdf(uploaded_file)
            
            if raw_text.startswith("ERROR"):
                status.update(label="❌ Lỗi đọc file!", state="error")
                st.stop()

            st.write("🔍 Đang tách danh mục và trích dẫn...")
            matches = list(re.finditer(r"(tài liệu tham khảo|references)", raw_text, re.IGNORECASE))
            if not matches:
                ref_raw = raw_text
                body_raw = raw_text
                st.toast("⚠️ Không tìm thấy tiêu đề References, quét toàn bộ.", icon="⚠️")
            else:
                split_idx = matches[-1].end()
                body_raw = raw_text[:matches[-1].start()]
                ref_raw = raw_text[split_idx:]
            
            body_text = preprocess_text(body_raw)
            ref_lines = [line.strip() for line in ref_raw.split('\n') if len(line.strip()) > 10 and re.search(r'\d{4}', line)]

            st.write("🧠 Đang chạy thuật toán AI Fuzzy Matching...")
            citations = find_citations_v12(body_text)

            # --- LOGIC CHECK (FUZZY) ---
            missing_refs = []
            for cit in citations:
                if not check_citation_fuzzy(cit['name'], cit['year'], ref_lines):
                    missing_refs.append(cit['full'])

            unused_refs = []
            for ref in ref_lines:
                if is_legal_or_standard(ref): continue
                
                ref_year_match = re.search(r'\d{4}', ref)
                if ref_year_match:
                    r_year = ref_year_match.group(0)
                    same_year_cites = [c for c in citations if c['year'] == r_year]
                    
                    is_found = False
                    if same_year_cites:
                        for c in same_year_cites:
                            if check_citation_fuzzy(c['name'], c['year'], [ref]):
                                is_found = True
                                break
                    if not is_found:
                        unused_refs.append(ref)
            
            status.update(label="✅ Đã phân tích xong!", state="complete", expanded=False)

    # --- DASHBOARD KẾT QUẢ ---
    
    st.markdown("<h3 style='margin-top: 20px;'>📊 Tổng quan (Dashboard)</h3>", unsafe_allow_html=True)
    
    st.markdown("""
    <div style="background-color: #ffe6e6; border: 1px solid #ffcccc; padding: 10px; border-radius: 5px; color: #cc0000; margin-bottom: 15px; font-size: 14px;">
        <b>⚠️ LƯU Ý:</b> Những trích dẫn bị xuống dòng trong bản thảo (ví dụ <i>Rasmussen</i> thành <i>Ras-mussen</i>) có thể bị báo lỗi thiếu trích dẫn do hạn chế của việc trích xuất văn bản PDF. Vui lòng kiểm tra lại thủ công.
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown('<p class="beta-note">(*) Kết quả dựa trên AI Fuzzy Logic. Vui lòng kiểm tra lại thủ công các mục báo lỗi.</p>', unsafe_allow_html=True)
    
    # Metrics
    m1, m2, m3, m4 = st.columns(4)
    with m1: st.metric("Tổng trích dẫn", len(citations), border=True)
    with m2: st.metric("Danh mục Ref", len(ref_lines), border=True)
    
    err_missing = len(missing_refs)
    err_unused = len(unused_refs)
    
    with m3: 
        st.metric("Lỗi thiếu Ref", err_missing, delta="-{}".format(err_missing) if err_missing > 0 else "OK", delta_color="inverse", border=True)
    with m4:
        st.metric("Lỗi thừa Ref", err_unused, delta="-{}".format(err_unused) if err_unused > 0 else "OK", delta_color="inverse", border=True)

    st.write("") 

    # Tabs
    tab_miss, tab_unused, tab_data = st.tabs(["🚫 TRÍCH DẪN THIẾU (Missing)", "⚠️ DANH MỤC THỪA (Unused)", "📋 DỮ LIỆU CHI TIẾT"])

    with tab_miss:
        st.markdown(f"**Danh sách {len(missing_refs)} trích dẫn có trong bài nhưng không tìm thấy trong danh mục:**")
        if missing_refs:
            for item in missing_refs:
                st.markdown(f'<div class="alert-error">❌ <b>{item}</b> - <i>Không tìm thấy nguồn</i></div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="alert-success">🎉 Tuyệt vời! Không có trích dẫn nào bị thiếu.</div>', unsafe_allow_html=True)

    with tab_unused:
        st.markdown(f"**Danh sách {len(unused_refs)} tài liệu có trong danh mục nhưng chưa được trích dẫn:**")
        if unused_refs:
            for item in unused_refs:
                st.markdown(f'<div class="alert-warning">⚠️ {item}</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="alert-success">🎉 Danh mục tài liệu rất gọn gàng.</div>', unsafe_allow_html=True)

    with tab_data:
        st.markdown("#### Tra cứu dữ liệu gốc")
        col_d1, col_d2 = st.columns(2)
        
        with col_d1:
            st.caption("Dữ liệu Trích dẫn (In-text)")
            if citations:
                df_cit = pd.DataFrame(citations)
                st.dataframe(df_cit, use_container_width=True, hide_index=True)
            else: st.info("Không có dữ liệu")

        with col_d2:
            st.caption("Dữ liệu Danh mục (References)")
            if ref_lines:
                df_ref = pd.DataFrame(ref_lines, columns=["Nội dung tham khảo"])
                st.dataframe(df_ref, use_container_width=True, hide_index=True)
            else: st.info("Không có dữ liệu")
